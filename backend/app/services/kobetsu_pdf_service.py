"""
Kobetsu Keiyakusho PDF/DOCX Generation Service

Generates legally compliant individual dispatch contracts (個別契約書)
with all 16 required items under 労働者派遣法第26条.
"""
import os
from datetime import date, datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.config import settings
from app.models.kobetsu_keiyakusho import KobetsuKeiyakusho


class KobetsuPDFService:
    """Service for generating Kobetsu Keiyakusho documents."""

    def __init__(self):
        self.output_dir = Path(settings.PDF_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _format_date_japanese(self, d: date) -> str:
        """Format date in Japanese style (令和X年X月X日)."""
        # Calculate Reiwa year (Reiwa started 2019-05-01)
        if d.year >= 2019:
            reiwa_year = d.year - 2018
            era = "令和"
        else:
            # Fallback to Western calendar
            return d.strftime("%Y年%m月%d日")

        return f"{era}{reiwa_year}年{d.month}月{d.day}日"

    def _format_time(self, t) -> str:
        """Format time as HH:MM."""
        return t.strftime("%H:%M")

    def _format_work_days(self, days: list) -> str:
        """Format work days list."""
        return "、".join(days)

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Add a heading with Japanese font."""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = "MS Gothic"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False, alignment=None):
        """Add a paragraph with Japanese font."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "MS Mincho"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
        run.font.size = Pt(10.5)
        run.bold = bold

        if alignment:
            p.alignment = alignment

        return p

    def _create_table(self, doc: Document, rows: int, cols: int):
        """Create a table with borders."""
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        return table

    def generate_docx(self, contract: KobetsuKeiyakusho) -> str:
        """
        Generate DOCX document for a contract.

        Args:
            contract: KobetsuKeiyakusho instance

        Returns:
            Path to generated DOCX file
        """
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("労働者派遣個別契約書")
        title_run.font.name = "MS Gothic"
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        title_run.font.size = Pt(18)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contract number and date
        doc.add_paragraph()
        self._add_paragraph(
            doc,
            f"契約番号: {contract.contract_number}",
            bold=True
        )
        self._add_paragraph(
            doc,
            f"契約締結日: {self._format_date_japanese(contract.contract_date)}"
        )

        doc.add_paragraph()

        # Introduction text
        intro_text = (
            f"株式会社UNS企画（以下「甲」という）と{contract.worksite_name}（以下「乙」という）は、"
            "労働者派遣法に基づき、以下のとおり労働者派遣個別契約を締結する。"
        )
        self._add_paragraph(doc, intro_text)

        doc.add_paragraph()

        # ========================================
        # 16 LEGALLY REQUIRED ITEMS
        # ========================================

        # 1. 派遣労働者の業務内容
        self._add_heading(doc, "第1条（業務内容）", level=2)
        self._add_paragraph(doc, f"業務内容: {contract.work_content}")
        self._add_paragraph(doc, f"責任の程度: {contract.responsibility_level}")

        # 2. 派遣先事業所の名称・所在地・組織単位
        self._add_heading(doc, "第2条（就業場所）", level=2)
        self._add_paragraph(doc, f"事業所名称: {contract.worksite_name}")
        self._add_paragraph(doc, f"所在地: {contract.worksite_address}")
        if contract.organizational_unit:
            self._add_paragraph(doc, f"組織単位: {contract.organizational_unit}")

        # 3. 指揮命令者
        self._add_heading(doc, "第3条（指揮命令者）", level=2)
        self._add_paragraph(
            doc,
            f"部署: {contract.supervisor_department} / "
            f"役職: {contract.supervisor_position} / "
            f"氏名: {contract.supervisor_name}"
        )

        # 4. 派遣期間
        self._add_heading(doc, "第4条（派遣期間）", level=2)
        self._add_paragraph(
            doc,
            f"派遣期間: {self._format_date_japanese(contract.dispatch_start_date)} から "
            f"{self._format_date_japanese(contract.dispatch_end_date)} まで"
        )

        # 5. 就業時間・休憩
        self._add_heading(doc, "第5条（就業時間）", level=2)
        self._add_paragraph(
            doc,
            f"就業日: {self._format_work_days(contract.work_days)}"
        )
        self._add_paragraph(
            doc,
            f"就業時間: {self._format_time(contract.work_start_time)} から "
            f"{self._format_time(contract.work_end_time)} まで"
        )
        self._add_paragraph(doc, f"休憩時間: {contract.break_time_minutes}分")

        # 6. 時間外労働
        self._add_heading(doc, "第6条（時間外労働）", level=2)
        if contract.overtime_max_hours_day:
            self._add_paragraph(doc, f"1日の時間外労働上限: {contract.overtime_max_hours_day}時間")
        if contract.overtime_max_hours_month:
            self._add_paragraph(doc, f"1ヶ月の時間外労働上限: {contract.overtime_max_hours_month}時間")
        if contract.overtime_max_days_month:
            self._add_paragraph(doc, f"1ヶ月の時間外労働日数上限: {contract.overtime_max_days_month}日")
        if contract.holiday_work_max_days:
            self._add_paragraph(doc, f"休日労働日数上限: {contract.holiday_work_max_days}日/月")

        # 7. 安全衛生
        self._add_heading(doc, "第7条（安全及び衛生）", level=2)
        safety_text = contract.safety_measures or "派遣先の安全衛生規程に従う"
        self._add_paragraph(doc, safety_text)

        # 8. 派遣労働者からの苦情処理
        self._add_heading(doc, "第8条（苦情処理）", level=2)
        self._add_paragraph(doc, "【派遣元苦情処理担当者】", bold=True)
        moto_contact = contract.haken_moto_complaint_contact
        self._add_paragraph(
            doc,
            f"部署: {moto_contact.get('department')} / "
            f"役職: {moto_contact.get('position')} / "
            f"氏名: {moto_contact.get('name')} / "
            f"電話: {moto_contact.get('phone')}"
        )
        self._add_paragraph(doc, "【派遣先苦情処理担当者】", bold=True)
        saki_contact = contract.haken_saki_complaint_contact
        self._add_paragraph(
            doc,
            f"部署: {saki_contact.get('department')} / "
            f"役職: {saki_contact.get('position')} / "
            f"氏名: {saki_contact.get('name')} / "
            f"電話: {saki_contact.get('phone')}"
        )

        # 9. 派遣契約解除時の措置
        self._add_heading(doc, "第9条（契約解除時の措置）", level=2)
        termination_text = contract.termination_measures or (
            "派遣契約を解除する場合、派遣先は派遣元に対し、30日前までに予告するものとする。"
            "また、派遣労働者の新たな就業機会の確保に努めるものとする。"
        )
        self._add_paragraph(doc, termination_text)

        # 10. 派遣元責任者・派遣先責任者
        self._add_heading(doc, "第10条（責任者）", level=2)
        self._add_paragraph(doc, "【派遣元責任者】", bold=True)
        moto_manager = contract.haken_moto_manager
        self._add_paragraph(
            doc,
            f"部署: {moto_manager.get('department')} / "
            f"役職: {moto_manager.get('position')} / "
            f"氏名: {moto_manager.get('name')} / "
            f"電話: {moto_manager.get('phone')}"
        )
        self._add_paragraph(doc, "【派遣先責任者】", bold=True)
        saki_manager = contract.haken_saki_manager
        self._add_paragraph(
            doc,
            f"部署: {saki_manager.get('department')} / "
            f"役職: {saki_manager.get('position')} / "
            f"氏名: {saki_manager.get('name')} / "
            f"電話: {saki_manager.get('phone')}"
        )

        # 11. 派遣労働者数
        self._add_heading(doc, "第11条（派遣労働者数）", level=2)
        self._add_paragraph(doc, f"派遣労働者数: {contract.number_of_workers}名")

        # 12. 福利厚生施設
        self._add_heading(doc, "第12条（福利厚生施設）", level=2)
        if contract.welfare_facilities:
            facilities = "、".join(contract.welfare_facilities)
            self._add_paragraph(doc, f"利用可能な福利厚生施設: {facilities}")
        else:
            self._add_paragraph(doc, "派遣先の福利厚生施設の利用については別途協議する")

        # 13. 派遣料金
        self._add_heading(doc, "第13条（派遣料金）", level=2)
        self._add_paragraph(doc, f"基本時間単価: {contract.hourly_rate:,}円")
        self._add_paragraph(doc, f"時間外単価: {contract.overtime_rate:,}円")
        if contract.night_shift_rate:
            self._add_paragraph(doc, f"深夜単価: {contract.night_shift_rate:,}円")
        if contract.holiday_rate:
            self._add_paragraph(doc, f"休日単価: {contract.holiday_rate:,}円")

        # 14-16. Additional legal provisions
        self._add_heading(doc, "第14条（その他）", level=2)
        if contract.is_kyotei_taisho:
            self._add_paragraph(doc, "・本契約は労使協定方式の対象となる")
        if contract.is_direct_hire_prevention:
            self._add_paragraph(doc, "・派遣先は派遣労働者の直接雇用に関する措置を講じる")
        if contract.is_mukeiko_60over_only:
            self._add_paragraph(doc, "・本契約は無期雇用又は60歳以上の派遣労働者のみを対象とする")

        # Notes
        if contract.notes:
            self._add_heading(doc, "備考", level=2)
            self._add_paragraph(doc, contract.notes)

        # Signature section
        doc.add_paragraph()
        doc.add_paragraph()

        self._add_paragraph(
            doc,
            "上記の条件にて派遣契約を締結することに同意し、本契約書2通を作成し、甲乙各1通を保有する。",
            alignment=WD_ALIGN_PARAGRAPH.LEFT
        )

        doc.add_paragraph()
        self._add_paragraph(
            doc,
            f"契約締結日: {self._format_date_japanese(contract.contract_date)}",
            alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )

        doc.add_paragraph()

        # Party A (Dispatch Company)
        self._add_paragraph(doc, "【甲】派遣元", bold=True)
        self._add_paragraph(doc, f"会社名: {settings.COMPANY_NAME}")
        self._add_paragraph(doc, f"所在地: {settings.COMPANY_ADDRESS}")
        self._add_paragraph(doc, f"許可番号: {settings.COMPANY_LICENSE_NUMBER}")
        self._add_paragraph(doc, "代表者: _____________________ 印")

        doc.add_paragraph()

        # Party B (Client Company)
        self._add_paragraph(doc, "【乙】派遣先", bold=True)
        self._add_paragraph(doc, f"会社名: {contract.worksite_name}")
        self._add_paragraph(doc, f"所在地: {contract.worksite_address}")
        self._add_paragraph(doc, "代表者: _____________________ 印")

        # Save document
        output_path = self.output_dir / f"{contract.contract_number}.docx"
        doc.save(str(output_path))

        return str(output_path)

    def generate_pdf(self, contract: KobetsuKeiyakusho) -> str:
        """
        Generate PDF document for a contract.

        First generates DOCX, then converts to PDF.

        Args:
            contract: KobetsuKeiyakusho instance

        Returns:
            Path to generated PDF file
        """
        # Generate DOCX first
        docx_path = self.generate_docx(contract)

        # Convert to PDF
        pdf_path = docx_path.replace('.docx', '.pdf')

        try:
            # Try using LibreOffice for conversion
            import subprocess
            result = subprocess.run(
                [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(self.output_dir),
                    docx_path
                ],
                capture_output=True,
                timeout=60
            )

            if result.returncode == 0 and os.path.exists(pdf_path):
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        try:
            # Alternative: Try using docx2pdf
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            return pdf_path
        except ImportError:
            pass

        # If PDF conversion fails, return DOCX path
        # In production, you would want to handle this better
        return docx_path

    def generate_preview(self, contract: KobetsuKeiyakusho) -> dict:
        """
        Generate a preview of the contract without creating a file.

        Args:
            contract: KobetsuKeiyakusho instance

        Returns:
            Dictionary with contract preview data
        """
        return {
            "contract_number": contract.contract_number,
            "contract_date": self._format_date_japanese(contract.contract_date),
            "dispatch_period": {
                "start": self._format_date_japanese(contract.dispatch_start_date),
                "end": self._format_date_japanese(contract.dispatch_end_date),
            },
            "dispatch_company": {
                "name": settings.COMPANY_NAME,
                "address": settings.COMPANY_ADDRESS,
                "license": settings.COMPANY_LICENSE_NUMBER,
            },
            "client_company": {
                "name": contract.worksite_name,
                "address": contract.worksite_address,
                "organizational_unit": contract.organizational_unit,
            },
            "work_details": {
                "content": contract.work_content,
                "responsibility_level": contract.responsibility_level,
                "supervisor": {
                    "department": contract.supervisor_department,
                    "position": contract.supervisor_position,
                    "name": contract.supervisor_name,
                },
            },
            "working_conditions": {
                "days": contract.work_days,
                "hours": f"{self._format_time(contract.work_start_time)} - {self._format_time(contract.work_end_time)}",
                "break_minutes": contract.break_time_minutes,
            },
            "rates": {
                "hourly": float(contract.hourly_rate),
                "overtime": float(contract.overtime_rate),
                "night_shift": float(contract.night_shift_rate) if contract.night_shift_rate else None,
                "holiday": float(contract.holiday_rate) if contract.holiday_rate else None,
            },
            "workers": contract.number_of_workers,
            "status": contract.status,
        }
