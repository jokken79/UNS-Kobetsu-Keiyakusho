"""
Dispatch Document Generation Service - 派遣書類生成サービス

Generates all legally required dispatch documents in compact A4 format:
1. 個別契約書 (Individual Contract) - 1 page
2. 就業条件明示書 (Working Conditions) - 1 page
3. 派遣通知書 (Dispatch Notification) - 1 page
4. 派遣先管理台帳 (Destination Ledger) - 1 page
5. 派遣元管理台帳 (Source Ledger) - 1 page
"""
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Cm, Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.config import settings


class DispatchDocumentService:
    """Service for generating all dispatch-related documents."""

    def __init__(self):
        self.output_dir = Path(settings.PDF_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ========================================
    # UTILITY METHODS
    # ========================================

    def _format_date_japanese(self, d: date) -> str:
        """Format date in Japanese style (令和X年X月X日)."""
        if d is None:
            return ""
        if d.year >= 2019:
            reiwa_year = d.year - 2018
            return f"令和{reiwa_year}年{d.month}月{d.day}日"
        return d.strftime("%Y年%m月%d日")

    def _format_date_short(self, d: date) -> str:
        """Short date format (R6.12.1)."""
        if d is None:
            return ""
        if d.year >= 2019:
            reiwa_year = d.year - 2018
            return f"R{reiwa_year}.{d.month}.{d.day}"
        return d.strftime("%Y.%m.%d")

    def _format_time(self, t) -> str:
        """Format time as HH:MM."""
        if t is None:
            return ""
        return t.strftime("%H:%M")

    def _setup_document_a4(self) -> Document:
        """Create a new A4 document with minimal margins."""
        doc = Document()
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(12)
        section.right_margin = Mm(12)
        return doc

    def _set_cell_font(self, cell, text: str, size: int = 8, bold: bool = False):
        """Set cell text with Japanese font."""
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(size)
        run.bold = bold
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _add_title(self, doc: Document, text: str, size: int = 14):
        """Add centered title."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(size)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _create_bordered_table(self, doc: Document, rows: int, cols: int):
        """Create table with borders."""
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        return table

    # ========================================
    # 1. 個別契約書 (COMPACT - 1 PAGE A4)
    # ========================================

    def generate_kobetsu_keiyakusho(self, data: Dict[str, Any]) -> bytes:
        """
        Generate compact 個別契約書 (Individual Contract) - 1 page A4.

        Required data keys:
        - contract_number, contract_date
        - dispatch_start_date, dispatch_end_date
        - client_company_name, client_address
        - worksite_name, worksite_address, organizational_unit
        - work_content, responsibility_level
        - supervisor_dept, supervisor_position, supervisor_name
        - work_days, work_start_time, work_end_time, break_minutes
        - overtime_max_day, overtime_max_month
        - hourly_rate, overtime_rate, holiday_rate
        - number_of_workers
        - haken_moto_manager, haken_saki_manager
        - haken_moto_complaint, haken_saki_complaint
        """
        doc = self._setup_document_a4()

        # Title
        self._add_title(doc, "労働者派遣個別契約書")

        # Contract info line
        p = doc.add_paragraph()
        run = p.add_run(f"契約番号: {data.get('contract_number', '')}　　契約日: {self._format_date_japanese(data.get('contract_date'))}")
        run.font.size = Pt(9)
        run.font.name = "MS Gothic"

        # Main table - 16 items in compact format
        table = self._create_bordered_table(doc, 20, 4)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Mm(25)
            row.cells[1].width = Mm(60)
            row.cells[2].width = Mm(25)
            row.cells[3].width = Mm(60)

        # Row 0: Contract parties
        self._set_cell_font(table.rows[0].cells[0], "派遣元", 8, True)
        self._set_cell_font(table.rows[0].cells[1], settings.COMPANY_NAME, 8)
        self._set_cell_font(table.rows[0].cells[2], "派遣先", 8, True)
        self._set_cell_font(table.rows[0].cells[3], data.get('client_company_name', ''), 8)

        # Row 1: Period
        self._set_cell_font(table.rows[1].cells[0], "派遣期間", 8, True)
        period = f"{self._format_date_short(data.get('dispatch_start_date'))} ～ {self._format_date_short(data.get('dispatch_end_date'))}"
        self._set_cell_font(table.rows[1].cells[1], period, 8)
        self._set_cell_font(table.rows[1].cells[2], "派遣人数", 8, True)
        self._set_cell_font(table.rows[1].cells[3], f"{data.get('number_of_workers', 1)}名", 8)

        # Row 2: Worksite
        self._set_cell_font(table.rows[2].cells[0], "就業場所", 8, True)
        worksite = f"{data.get('worksite_name', '')}\n{data.get('worksite_address', '')}"
        self._set_cell_font(table.rows[2].cells[1], worksite, 7)
        self._set_cell_font(table.rows[2].cells[2], "組織単位", 8, True)
        self._set_cell_font(table.rows[2].cells[3], data.get('organizational_unit', ''), 8)

        # Row 3: Work content
        self._set_cell_font(table.rows[3].cells[0], "業務内容", 8, True)
        table.rows[3].cells[1].merge(table.rows[3].cells[3])
        self._set_cell_font(table.rows[3].cells[1], data.get('work_content', ''), 8)

        # Row 4: Responsibility
        self._set_cell_font(table.rows[4].cells[0], "責任の程度", 8, True)
        self._set_cell_font(table.rows[4].cells[1], data.get('responsibility_level', '通常業務'), 8)
        self._set_cell_font(table.rows[4].cells[2], "業務の種類", 8, True)
        self._set_cell_font(table.rows[4].cells[3], data.get('business_type', '製造業務'), 8)

        # Row 5: Supervisor
        self._set_cell_font(table.rows[5].cells[0], "指揮命令者", 8, True)
        supervisor = f"{data.get('supervisor_dept', '')} {data.get('supervisor_position', '')} {data.get('supervisor_name', '')}"
        table.rows[5].cells[1].merge(table.rows[5].cells[3])
        self._set_cell_font(table.rows[5].cells[1], supervisor, 8)

        # Row 6: Working hours
        self._set_cell_font(table.rows[6].cells[0], "就業日", 8, True)
        work_days = data.get('work_days', ['月', '火', '水', '木', '金'])
        if isinstance(work_days, list):
            work_days = '・'.join(work_days)
        self._set_cell_font(table.rows[6].cells[1], work_days, 8)
        self._set_cell_font(table.rows[6].cells[2], "休憩", 8, True)
        self._set_cell_font(table.rows[6].cells[3], f"{data.get('break_minutes', 60)}分", 8)

        # Row 7: Time
        self._set_cell_font(table.rows[7].cells[0], "就業時間", 8, True)
        time_str = f"{self._format_time(data.get('work_start_time'))} ～ {self._format_time(data.get('work_end_time'))}"
        self._set_cell_font(table.rows[7].cells[1], time_str, 8)
        self._set_cell_font(table.rows[7].cells[2], "実働", 8, True)
        self._set_cell_font(table.rows[7].cells[3], f"{data.get('actual_hours', 8)}時間", 8)

        # Row 8: Overtime
        self._set_cell_font(table.rows[8].cells[0], "時間外労働", 8, True)
        overtime = f"1日{data.get('overtime_max_day', 4)}H以内 / 月{data.get('overtime_max_month', 45)}H以内"
        self._set_cell_font(table.rows[8].cells[1], overtime, 8)
        self._set_cell_font(table.rows[8].cells[2], "休日労働", 8, True)
        self._set_cell_font(table.rows[8].cells[3], f"月{data.get('holiday_work_max', 2)}日以内", 8)

        # Row 9: Rates
        self._set_cell_font(table.rows[9].cells[0], "派遣料金", 8, True)
        rate = f"基本: ¥{data.get('hourly_rate', 0):,.0f}/H"
        self._set_cell_font(table.rows[9].cells[1], rate, 8)
        self._set_cell_font(table.rows[9].cells[2], "時間外等", 8, True)
        ot_rate = f"残業: ¥{data.get('overtime_rate', 0):,.0f} 休日: ¥{data.get('holiday_rate', 0):,.0f}"
        self._set_cell_font(table.rows[9].cells[3], ot_rate, 7)

        # Row 10: Haken-moto manager
        self._set_cell_font(table.rows[10].cells[0], "派遣元責任者", 8, True)
        moto_mgr = data.get('haken_moto_manager', {})
        self._set_cell_font(table.rows[10].cells[1], f"{moto_mgr.get('name', '')} TEL:{moto_mgr.get('phone', '')}", 7)
        self._set_cell_font(table.rows[10].cells[2], "派遣先責任者", 8, True)
        saki_mgr = data.get('haken_saki_manager', {})
        self._set_cell_font(table.rows[10].cells[3], f"{saki_mgr.get('name', '')} TEL:{saki_mgr.get('phone', '')}", 7)

        # Row 11: Complaint handlers
        self._set_cell_font(table.rows[11].cells[0], "派遣元苦情処理", 8, True)
        moto_comp = data.get('haken_moto_complaint', {})
        self._set_cell_font(table.rows[11].cells[1], f"{moto_comp.get('name', '')} TEL:{moto_comp.get('phone', '')}", 7)
        self._set_cell_font(table.rows[11].cells[2], "派遣先苦情処理", 8, True)
        saki_comp = data.get('haken_saki_complaint', {})
        self._set_cell_font(table.rows[11].cells[3], f"{saki_comp.get('name', '')} TEL:{saki_comp.get('phone', '')}", 7)

        # Row 12: Safety
        self._set_cell_font(table.rows[12].cells[0], "安全衛生", 8, True)
        table.rows[12].cells[1].merge(table.rows[12].cells[3])
        self._set_cell_font(table.rows[12].cells[1], data.get('safety_measures', '派遣先の安全衛生規程に従う'), 7)

        # Row 13: Termination measures
        self._set_cell_font(table.rows[13].cells[0], "契約解除措置", 8, True)
        table.rows[13].cells[1].merge(table.rows[13].cells[3])
        self._set_cell_font(table.rows[13].cells[1], data.get('termination_measures', '30日前予告。派遣労働者の就業機会確保に努める。'), 7)

        # Row 14: Welfare
        self._set_cell_font(table.rows[14].cells[0], "福利厚生施設", 8, True)
        welfare = data.get('welfare_facilities', ['食堂', '更衣室', '休憩室'])
        if isinstance(welfare, list):
            welfare = '、'.join(welfare)
        table.rows[14].cells[1].merge(table.rows[14].cells[3])
        self._set_cell_font(table.rows[14].cells[1], welfare, 8)

        # Row 15: Legal checkboxes
        self._set_cell_font(table.rows[15].cells[0], "労使協定方式", 8, True)
        kyotei = "☑ 対象" if data.get('is_kyotei_taisho') else "☐ 対象外"
        self._set_cell_font(table.rows[15].cells[1], kyotei, 8)
        self._set_cell_font(table.rows[15].cells[2], "無期/60歳以上", 8, True)
        mukeiko = "☑ 該当" if data.get('is_mukeiko_60over') else "☐ 非該当"
        self._set_cell_font(table.rows[15].cells[3], mukeiko, 8)

        # Row 16-17: Signatures (派遣元)
        self._set_cell_font(table.rows[16].cells[0], "【甲】派遣元", 8, True)
        table.rows[16].cells[1].merge(table.rows[16].cells[3])
        self._set_cell_font(table.rows[16].cells[1], f"{settings.COMPANY_NAME}　許可番号: {settings.COMPANY_LICENSE_NUMBER}", 8)

        self._set_cell_font(table.rows[17].cells[0], "住所", 8)
        table.rows[17].cells[1].merge(table.rows[17].cells[3])
        self._set_cell_font(table.rows[17].cells[1], f"{settings.COMPANY_ADDRESS}　　　　　　　　　　　　　　印", 8)

        # Row 18-19: Signatures (派遣先)
        self._set_cell_font(table.rows[18].cells[0], "【乙】派遣先", 8, True)
        table.rows[18].cells[1].merge(table.rows[18].cells[3])
        self._set_cell_font(table.rows[18].cells[1], data.get('client_company_name', ''), 8)

        self._set_cell_font(table.rows[19].cells[0], "住所", 8)
        table.rows[19].cells[1].merge(table.rows[19].cells[3])
        self._set_cell_font(table.rows[19].cells[1], f"{data.get('client_address', '')}　　　　　　　　　　　　　　印", 8)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ========================================
    # 2. 就業条件明示書 (1 PAGE A4)
    # ========================================

    def generate_shugyo_joken_meijisho(self, data: Dict[str, Any]) -> bytes:
        """
        Generate 就業条件明示書 (Working Conditions Document) - 1 page A4.
        派遣元 → 派遣労働者

        Required data keys:
        - worker_name, worker_number
        - dispatch_start_date, dispatch_end_date
        - worksite_name, worksite_address
        - work_content, responsibility_level
        - work_days, work_start_time, work_end_time, break_minutes
        - hourly_wage, overtime_wage
        - haken_moto_manager, complaint_handler
        - conflict_date (抵触日)
        """
        doc = self._setup_document_a4()

        # Title
        self._add_title(doc, "就業条件明示書")

        # Issue date and recipient
        p = doc.add_paragraph()
        run = p.add_run(f"交付日: {self._format_date_japanese(date.today())}　　　{data.get('worker_name', '')} 殿")
        run.font.size = Pt(10)
        run.font.name = "MS Gothic"

        # Issuer info
        p = doc.add_paragraph()
        run = p.add_run(f"派遣元事業主: {settings.COMPANY_NAME}　許可番号: {settings.COMPANY_LICENSE_NUMBER}")
        run.font.size = Pt(9)
        run.font.name = "MS Gothic"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Main table
        table = self._create_bordered_table(doc, 18, 4)

        # Column widths
        for row in table.rows:
            row.cells[0].width = Mm(30)
            row.cells[1].width = Mm(55)
            row.cells[2].width = Mm(30)
            row.cells[3].width = Mm(55)

        row_idx = 0

        # Row 0: Contract period
        self._set_cell_font(table.rows[row_idx].cells[0], "雇用契約期間", 8, True)
        period = f"{self._format_date_short(data.get('dispatch_start_date'))} ～ {self._format_date_short(data.get('dispatch_end_date'))}"
        self._set_cell_font(table.rows[row_idx].cells[1], period, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "更新の有無", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('renewal_policy', '更新する場合がある'), 8)

        row_idx += 1
        # Row 1: Worksite
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣先名称", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('worksite_name', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "派遣先所在地", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('worksite_address', ''), 7)

        row_idx += 1
        # Row 2: Organizational unit
        self._set_cell_font(table.rows[row_idx].cells[0], "組織単位", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('organizational_unit', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "指揮命令者", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('supervisor_name', ''), 8)

        row_idx += 1
        # Row 3: Work content
        self._set_cell_font(table.rows[row_idx].cells[0], "業務内容", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('work_content', ''), 8)

        row_idx += 1
        # Row 4: Responsibility
        self._set_cell_font(table.rows[row_idx].cells[0], "責任の程度", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('responsibility_level', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "業務変更範囲", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('work_change_scope', '会社の定める業務'), 7)

        row_idx += 1
        # Row 5: Work days
        self._set_cell_font(table.rows[row_idx].cells[0], "就業日", 8, True)
        work_days = data.get('work_days', [])
        if isinstance(work_days, list):
            work_days = '・'.join(work_days)
        self._set_cell_font(table.rows[row_idx].cells[1], work_days, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "休日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('holidays', '土・日・祝日'), 8)

        row_idx += 1
        # Row 6: Work time
        self._set_cell_font(table.rows[row_idx].cells[0], "就業時間", 8, True)
        time_str = f"{self._format_time(data.get('work_start_time'))} ～ {self._format_time(data.get('work_end_time'))}"
        self._set_cell_font(table.rows[row_idx].cells[1], time_str, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "休憩時間", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], f"{data.get('break_minutes', 60)}分", 8)

        row_idx += 1
        # Row 7: Overtime
        self._set_cell_font(table.rows[row_idx].cells[0], "時間外労働", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], f"有（月{data.get('overtime_max_month', 45)}時間以内）", 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "休日労働", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], f"有（月{data.get('holiday_work_max', 2)}日以内）", 8)

        row_idx += 1
        # Row 8: Wages
        self._set_cell_font(table.rows[row_idx].cells[0], "賃金", 8, True)
        wage = f"時給 ¥{data.get('hourly_wage', 0):,.0f}"
        self._set_cell_font(table.rows[row_idx].cells[1], wage, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "時間外等", 8, True)
        ot_wage = f"125%: ¥{data.get('overtime_wage', 0):,.0f}"
        self._set_cell_font(table.rows[row_idx].cells[3], ot_wage, 8)

        row_idx += 1
        # Row 9: Payment
        self._set_cell_font(table.rows[row_idx].cells[0], "賃金締切日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('wage_closing', '月末日'), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "賃金支払日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('wage_payment', '翌月25日'), 8)

        row_idx += 1
        # Row 10: Insurance
        self._set_cell_font(table.rows[row_idx].cells[0], "社会保険", 8, True)
        insurance = []
        if data.get('has_health_insurance', True):
            insurance.append("健康保険")
        if data.get('has_pension', True):
            insurance.append("厚生年金")
        if data.get('has_employment_insurance', True):
            insurance.append("雇用保険")
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], "・".join(insurance) + " 加入", 8)

        row_idx += 1
        # Row 11: Conflict date (抵触日)
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣先抵触日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], self._format_date_japanese(data.get('conflict_date')), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "個人抵触日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], self._format_date_japanese(data.get('personal_conflict_date')), 8)

        row_idx += 1
        # Row 12: Employment type
        self._set_cell_font(table.rows[row_idx].cells[0], "雇用形態", 8, True)
        emp_type = "☑無期雇用 ☐有期雇用" if data.get('is_indefinite') else "☐無期雇用 ☑有期雇用"
        self._set_cell_font(table.rows[row_idx].cells[1], emp_type, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "派遣労働者区分", 8, True)
        agreement = "☑協定対象" if data.get('is_agreement_target', True) else "☐協定対象外"
        self._set_cell_font(table.rows[row_idx].cells[3], agreement, 8)

        row_idx += 1
        # Row 13: Haken-moto manager
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣元責任者", 8, True)
        moto_mgr = data.get('haken_moto_manager', {})
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], f"{moto_mgr.get('name', '')}　TEL: {moto_mgr.get('phone', '')}", 8)

        row_idx += 1
        # Row 14: Complaint handler (派遣元)
        self._set_cell_font(table.rows[row_idx].cells[0], "苦情処理(派遣元)", 8, True)
        comp_moto = data.get('complaint_handler_moto', {})
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], f"{comp_moto.get('name', '')}　TEL: {comp_moto.get('phone', '')}", 8)

        row_idx += 1
        # Row 15: Complaint handler (派遣先)
        self._set_cell_font(table.rows[row_idx].cells[0], "苦情処理(派遣先)", 8, True)
        comp_saki = data.get('complaint_handler_saki', {})
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], f"{comp_saki.get('name', '')}　TEL: {comp_saki.get('phone', '')}", 8)

        row_idx += 1
        # Row 16: Welfare
        self._set_cell_font(table.rows[row_idx].cells[0], "福利厚生", 8, True)
        welfare = data.get('welfare_facilities', [])
        if isinstance(welfare, list):
            welfare = '、'.join(welfare)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], welfare, 8)

        row_idx += 1
        # Row 17: Notes
        self._set_cell_font(table.rows[row_idx].cells[0], "備考", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('notes', ''), 8)

        # Footer
        p = doc.add_paragraph()
        run = p.add_run("\n上記の就業条件を明示します。")
        run.font.size = Pt(9)
        run.font.name = "MS Gothic"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ========================================
    # 3. 派遣通知書 (1 PAGE A4)
    # ========================================

    def generate_haken_tsuchisho(self, data: Dict[str, Any]) -> bytes:
        """
        Generate 派遣通知書 (Dispatch Notification) - 1 page A4.
        派遣元 → 派遣先
        労働者派遣法第35条に基づく通知

        Required data keys:
        - worker_name, worker_gender
        - is_indefinite (無期/有期)
        - is_over_60
        - has_health_insurance, has_pension, has_employment_insurance
        - dispatch_start_date, dispatch_end_date
        - work_content
        """
        doc = self._setup_document_a4()

        # Title
        self._add_title(doc, "派遣労働者通知書")

        # Subtitle
        p = doc.add_paragraph()
        run = p.add_run("（労働者派遣法第35条第1項に基づく通知）")
        run.font.size = Pt(9)
        run.font.name = "MS Gothic"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date and recipient
        p = doc.add_paragraph()
        run = p.add_run(f"\n通知日: {self._format_date_japanese(date.today())}")
        run.font.size = Pt(10)
        run.font.name = "MS Gothic"

        p = doc.add_paragraph()
        run = p.add_run(f"{data.get('client_company_name', '')} 御中")
        run.font.size = Pt(11)
        run.font.name = "MS Gothic"

        # From
        p = doc.add_paragraph()
        run = p.add_run(f"派遣元事業主: {settings.COMPANY_NAME}")
        run.font.size = Pt(10)
        run.font.name = "MS Gothic"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = doc.add_paragraph()
        run = p.add_run(f"許可番号: {settings.COMPANY_LICENSE_NUMBER}")
        run.font.size = Pt(9)
        run.font.name = "MS Gothic"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Introduction
        p = doc.add_paragraph()
        run = p.add_run("\n下記の労働者を派遣いたしますので、労働者派遣法第35条第1項に基づき通知いたします。")
        run.font.size = Pt(10)
        run.font.name = "MS Gothic"

        # Worker table
        workers = data.get('workers', [data])  # Support single or multiple workers

        table = self._create_bordered_table(doc, len(workers) + 1, 8)

        # Header row
        headers = ['氏名', '性別', '無期/有期', '60歳以上', '健康保険', '厚生年金', '雇用保険', '協定対象']
        for i, header in enumerate(headers):
            self._set_cell_font(table.rows[0].cells[i], header, 7, True)

        # Data rows
        for idx, worker in enumerate(workers, 1):
            self._set_cell_font(table.rows[idx].cells[0], worker.get('worker_name', ''), 8)
            self._set_cell_font(table.rows[idx].cells[1], worker.get('worker_gender', ''), 8)
            emp_type = "無期" if worker.get('is_indefinite') else "有期"
            self._set_cell_font(table.rows[idx].cells[2], emp_type, 8)
            over_60 = "該当" if worker.get('is_over_60') else "非該当"
            self._set_cell_font(table.rows[idx].cells[3], over_60, 8)
            health = "加入" if worker.get('has_health_insurance', True) else "未加入"
            self._set_cell_font(table.rows[idx].cells[4], health, 8)
            pension = "加入" if worker.get('has_pension', True) else "未加入"
            self._set_cell_font(table.rows[idx].cells[5], pension, 8)
            emp_ins = "加入" if worker.get('has_employment_insurance', True) else "未加入"
            self._set_cell_font(table.rows[idx].cells[6], emp_ins, 8)
            agreement = "対象" if worker.get('is_agreement_target', True) else "対象外"
            self._set_cell_font(table.rows[idx].cells[7], agreement, 8)

        # Contract info table
        doc.add_paragraph()
        info_table = self._create_bordered_table(doc, 4, 4)

        self._set_cell_font(info_table.rows[0].cells[0], "派遣期間", 8, True)
        period = f"{self._format_date_short(data.get('dispatch_start_date'))} ～ {self._format_date_short(data.get('dispatch_end_date'))}"
        info_table.rows[0].cells[1].merge(info_table.rows[0].cells[3])
        self._set_cell_font(info_table.rows[0].cells[1], period, 8)

        self._set_cell_font(info_table.rows[1].cells[0], "就業場所", 8, True)
        info_table.rows[1].cells[1].merge(info_table.rows[1].cells[3])
        self._set_cell_font(info_table.rows[1].cells[1], data.get('worksite_name', ''), 8)

        self._set_cell_font(info_table.rows[2].cells[0], "業務内容", 8, True)
        info_table.rows[2].cells[1].merge(info_table.rows[2].cells[3])
        self._set_cell_font(info_table.rows[2].cells[1], data.get('work_content', ''), 8)

        self._set_cell_font(info_table.rows[3].cells[0], "派遣元責任者", 8, True)
        moto_mgr = data.get('haken_moto_manager', {})
        self._set_cell_font(info_table.rows[3].cells[1], moto_mgr.get('name', ''), 8)
        self._set_cell_font(info_table.rows[3].cells[2], "連絡先", 8, True)
        self._set_cell_font(info_table.rows[3].cells[3], moto_mgr.get('phone', ''), 8)

        # Note
        p = doc.add_paragraph()
        run = p.add_run("\n※本通知書は、派遣先管理台帳の作成・保管にご利用ください。")
        run.font.size = Pt(8)
        run.font.name = "MS Gothic"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ========================================
    # 4. 派遣先管理台帳 (1 PAGE A4)
    # ========================================

    def generate_hakensaki_kanri_daicho(self, data: Dict[str, Any]) -> bytes:
        """
        Generate 派遣先管理台帳 (Destination Management Ledger) - 1 page A4.
        派遣先が作成・保管する台帳
        """
        doc = self._setup_document_a4()

        # Title
        self._add_title(doc, "派遣先管理台帳")

        # Worker info
        p = doc.add_paragraph()
        run = p.add_run(f"派遣労働者氏名: {data.get('worker_name', '')}")
        run.font.size = Pt(11)
        run.font.name = "MS Gothic"

        # Main table
        table = self._create_bordered_table(doc, 16, 4)

        for row in table.rows:
            row.cells[0].width = Mm(30)
            row.cells[1].width = Mm(55)
            row.cells[2].width = Mm(30)
            row.cells[3].width = Mm(55)

        row_idx = 0

        # Row 0: Dispatch company
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣元事業主", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('dispatch_company', settings.COMPANY_NAME), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "事業所名", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('dispatch_office', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "事業所所在地", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('dispatch_address', settings.COMPANY_ADDRESS), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "業務内容", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('work_content', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "責任の程度", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('responsibility_level', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "協定対象", 8, True)
        agreement = "対象" if data.get('is_agreement_target', True) else "対象外"
        self._set_cell_font(table.rows[row_idx].cells[3], agreement, 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "無期/有期", 8, True)
        emp_type = "無期雇用" if data.get('is_indefinite') else "有期雇用"
        self._set_cell_font(table.rows[row_idx].cells[1], emp_type, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "60歳以上", 8, True)
        over_60 = "該当" if data.get('is_over_60') else "非該当"
        self._set_cell_font(table.rows[row_idx].cells[3], over_60, 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣期間", 8, True)
        period = f"{self._format_date_short(data.get('dispatch_start_date'))} ～ {self._format_date_short(data.get('dispatch_end_date'))}"
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], period, 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "就業場所", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], f"{data.get('worksite_name', '')} {data.get('organizational_unit', '')}", 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "組織単位", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('organizational_unit', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "指揮命令者", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('supervisor_name', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "就業日", 8, True)
        work_days = data.get('work_days', [])
        if isinstance(work_days, list):
            work_days = '・'.join(work_days)
        self._set_cell_font(table.rows[row_idx].cells[1], work_days, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "就業時間", 8, True)
        time_str = f"{self._format_time(data.get('work_start_time'))} ～ {self._format_time(data.get('work_end_time'))}"
        self._set_cell_font(table.rows[row_idx].cells[3], time_str, 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣元責任者", 8, True)
        moto_mgr = data.get('haken_moto_manager', {})
        self._set_cell_font(table.rows[row_idx].cells[1], moto_mgr.get('name', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "派遣先責任者", 8, True)
        saki_mgr = data.get('haken_saki_manager', {})
        self._set_cell_font(table.rows[row_idx].cells[3], saki_mgr.get('name', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "苦情処理(元)", 8, True)
        comp_moto = data.get('complaint_handler_moto', {})
        self._set_cell_font(table.rows[row_idx].cells[1], comp_moto.get('name', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "苦情処理(先)", 8, True)
        comp_saki = data.get('complaint_handler_saki', {})
        self._set_cell_font(table.rows[row_idx].cells[3], comp_saki.get('name', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "社会保険", 8, True)
        insurance = []
        if data.get('has_health_insurance', True):
            insurance.append("健保")
        if data.get('has_pension', True):
            insurance.append("年金")
        if data.get('has_employment_insurance', True):
            insurance.append("雇保")
        self._set_cell_font(table.rows[row_idx].cells[1], "・".join(insurance), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "教育訓練", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('training', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "福利厚生", 8, True)
        welfare = data.get('welfare_facilities', [])
        if isinstance(welfare, list):
            welfare = '、'.join(welfare)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], welfare, 8)

        # Work record section header
        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "【就業実績】", 8, True)
        table.rows[row_idx].cells[0].merge(table.rows[row_idx].cells[3])

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "就業日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], "始業・終業時刻", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[2], "休憩時間", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], "備考", 8, True)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "", 8)
        self._set_cell_font(table.rows[row_idx].cells[1], "", 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "", 8)
        self._set_cell_font(table.rows[row_idx].cells[3], "", 8)

        # Footer note
        p = doc.add_paragraph()
        run = p.add_run("\n※派遣終了後3年間保存要")
        run.font.size = Pt(8)
        run.font.name = "MS Gothic"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ========================================
    # 5. 派遣元管理台帳 (1 PAGE A4)
    # ========================================

    def generate_hakenmoto_kanri_daicho(self, data: Dict[str, Any]) -> bytes:
        """
        Generate 派遣元管理台帳 (Source Management Ledger) - 1 page A4.
        派遣元が作成・保管する台帳
        """
        doc = self._setup_document_a4()

        # Title
        self._add_title(doc, "派遣元管理台帳")

        # Worker basic info
        p = doc.add_paragraph()
        run = p.add_run(f"派遣労働者: {data.get('worker_name', '')} （社員番号: {data.get('worker_number', '')}）")
        run.font.size = Pt(11)
        run.font.name = "MS Gothic"

        # Main table
        table = self._create_bordered_table(doc, 18, 4)

        for row in table.rows:
            row.cells[0].width = Mm(30)
            row.cells[1].width = Mm(55)
            row.cells[2].width = Mm(30)
            row.cells[3].width = Mm(55)

        row_idx = 0

        # Personal info
        self._set_cell_font(table.rows[row_idx].cells[0], "氏名", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('worker_name', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "性別", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('worker_gender', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "雇用形態", 8, True)
        emp_type = "無期雇用" if data.get('is_indefinite') else "有期雇用"
        self._set_cell_font(table.rows[row_idx].cells[1], emp_type, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "60歳以上", 8, True)
        over_60 = "該当" if data.get('is_over_60') else "非該当"
        self._set_cell_font(table.rows[row_idx].cells[3], over_60, 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣先名称", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('client_company_name', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "事業所", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('worksite_name', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣先所在地", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('worksite_address', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "組織単位", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('organizational_unit', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "指揮命令者", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('supervisor_name', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "業務内容", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('work_content', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "責任の程度", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('responsibility_level', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "協定対象", 8, True)
        agreement = "対象" if data.get('is_agreement_target', True) else "対象外"
        self._set_cell_font(table.rows[row_idx].cells[3], agreement, 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣期間", 8, True)
        period = f"{self._format_date_short(data.get('dispatch_start_date'))} ～ {self._format_date_short(data.get('dispatch_end_date'))}"
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], period, 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "就業日", 8, True)
        work_days = data.get('work_days', [])
        if isinstance(work_days, list):
            work_days = '・'.join(work_days)
        self._set_cell_font(table.rows[row_idx].cells[1], work_days, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "就業時間", 8, True)
        time_str = f"{self._format_time(data.get('work_start_time'))} ～ {self._format_time(data.get('work_end_time'))}"
        self._set_cell_font(table.rows[row_idx].cells[3], time_str, 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣料金", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], f"¥{data.get('billing_rate', 0):,.0f}/H", 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "賃金", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], f"¥{data.get('hourly_wage', 0):,.0f}/H", 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣先責任者", 8, True)
        saki_mgr = data.get('haken_saki_manager', {})
        self._set_cell_font(table.rows[row_idx].cells[1], saki_mgr.get('name', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "連絡先", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], saki_mgr.get('phone', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "社会保険", 8, True)
        insurance = []
        if data.get('has_health_insurance', True):
            insurance.append("健保")
        if data.get('has_pension', True):
            insurance.append("年金")
        if data.get('has_employment_insurance', True):
            insurance.append("雇保")
        self._set_cell_font(table.rows[row_idx].cells[1], "・".join(insurance), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "教育訓練", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('training_date', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "キャリア相談", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('career_consultation_date', ''), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "雇用安定措置", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('employment_stability_measures', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "派遣先抵触日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], self._format_date_short(data.get('conflict_date')), 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "個人抵触日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], self._format_date_short(data.get('personal_conflict_date')), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "苦情申出日", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[1], "", 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "苦情内容", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], "", 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "苦情処理状況", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], "", 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "紹介予定派遣", 8, True)
        shokai = "該当" if data.get('is_temp_to_perm') else "非該当"
        self._set_cell_font(table.rows[row_idx].cells[1], shokai, 8)
        self._set_cell_font(table.rows[row_idx].cells[2], "直接雇用結果", 8, True)
        self._set_cell_font(table.rows[row_idx].cells[3], data.get('direct_hire_result', ''), 8)

        row_idx += 1
        self._set_cell_font(table.rows[row_idx].cells[0], "備考", 8, True)
        table.rows[row_idx].cells[1].merge(table.rows[row_idx].cells[3])
        self._set_cell_font(table.rows[row_idx].cells[1], data.get('notes', ''), 8)

        # Footer
        p = doc.add_paragraph()
        run = p.add_run("\n※派遣終了後3年間保存要　※苦情処理・教育訓練等は随時記録")
        run.font.size = Pt(8)
        run.font.name = "MS Gothic"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ========================================
    # 6. 個別契約書 + 就業条件明示書 一体型 (1 PAGE A4)
    # ========================================

    def generate_kobetsu_shugyo_combined(self, data: Dict[str, Any]) -> bytes:
        """
        Generate combined 個別契約書 + 就業条件明示書 on ONE A4 page.
        派遣元 ↔ 派遣先 契約 + 派遣労働者への条件明示
        Ultra-compact format with all legally required items.
        """
        doc = self._setup_document_a4()

        # Smaller margins for more space
        section = doc.sections[0]
        section.top_margin = Mm(8)
        section.bottom_margin = Mm(8)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)

        # Title
        p = doc.add_paragraph()
        run = p.add_run("労働者派遣個別契約書（兼）就業条件明示書")
        run.font.name = "MS Gothic"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contract number line
        p = doc.add_paragraph()
        run = p.add_run(f"契約番号: {data.get('contract_number', '')}　契約日: {self._format_date_short(data.get('contract_date'))}　派遣労働者: {data.get('worker_name', '')}")
        run.font.size = Pt(7)
        run.font.name = "MS Gothic"

        # Main table - very compact
        table = self._create_bordered_table(doc, 24, 6)

        # Column widths (total ~186mm usable)
        for row in table.rows:
            row.cells[0].width = Mm(18)
            row.cells[1].width = Mm(42)
            row.cells[2].width = Mm(18)
            row.cells[3].width = Mm(42)
            row.cells[4].width = Mm(18)
            row.cells[5].width = Mm(42)

        r = 0  # row index

        # Row 0: Parties
        self._set_cell_font(table.rows[r].cells[0], "派遣元", 6, True)
        self._set_cell_font(table.rows[r].cells[1], settings.COMPANY_NAME, 6)
        self._set_cell_font(table.rows[r].cells[2], "派遣先", 6, True)
        self._set_cell_font(table.rows[r].cells[3], data.get('client_company_name', ''), 6)
        self._set_cell_font(table.rows[r].cells[4], "許可番号", 6, True)
        self._set_cell_font(table.rows[r].cells[5], settings.COMPANY_LICENSE_NUMBER, 6)

        r += 1
        # Row 1: Period & Workers
        self._set_cell_font(table.rows[r].cells[0], "派遣期間", 6, True)
        period = f"{self._format_date_short(data.get('dispatch_start_date'))}～{self._format_date_short(data.get('dispatch_end_date'))}"
        self._set_cell_font(table.rows[r].cells[1], period, 6)
        self._set_cell_font(table.rows[r].cells[2], "派遣人数", 6, True)
        self._set_cell_font(table.rows[r].cells[3], f"{data.get('number_of_workers', 1)}名", 6)
        self._set_cell_font(table.rows[r].cells[4], "抵触日", 6, True)
        self._set_cell_font(table.rows[r].cells[5], self._format_date_short(data.get('conflict_date')), 6)

        r += 1
        # Row 2: Worksite
        self._set_cell_font(table.rows[r].cells[0], "就業場所", 6, True)
        table.rows[r].cells[1].merge(table.rows[r].cells[3])
        self._set_cell_font(table.rows[r].cells[1], data.get('worksite_name', ''), 6)
        self._set_cell_font(table.rows[r].cells[4], "組織単位", 6, True)
        self._set_cell_font(table.rows[r].cells[5], data.get('organizational_unit', ''), 6)

        r += 1
        # Row 3: Address
        self._set_cell_font(table.rows[r].cells[0], "所在地", 6, True)
        table.rows[r].cells[1].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[1], data.get('worksite_address', ''), 6)

        r += 1
        # Row 4: Work content
        self._set_cell_font(table.rows[r].cells[0], "業務内容", 6, True)
        table.rows[r].cells[1].merge(table.rows[r].cells[3])
        self._set_cell_font(table.rows[r].cells[1], data.get('work_content', ''), 6)
        self._set_cell_font(table.rows[r].cells[4], "責任程度", 6, True)
        self._set_cell_font(table.rows[r].cells[5], data.get('responsibility_level', ''), 6)

        r += 1
        # Row 5: Supervisor
        self._set_cell_font(table.rows[r].cells[0], "指揮命令者", 6, True)
        supervisor = f"{data.get('supervisor_dept', '')} {data.get('supervisor_name', '')}"
        table.rows[r].cells[1].merge(table.rows[r].cells[3])
        self._set_cell_font(table.rows[r].cells[1], supervisor, 6)
        self._set_cell_font(table.rows[r].cells[4], "TEL", 6, True)
        self._set_cell_font(table.rows[r].cells[5], data.get('supervisor_phone', ''), 6)

        r += 1
        # Row 6: Work days/time
        self._set_cell_font(table.rows[r].cells[0], "就業日", 6, True)
        work_days = data.get('work_days', [])
        if isinstance(work_days, list):
            work_days = '・'.join(work_days)
        self._set_cell_font(table.rows[r].cells[1], work_days, 6)
        self._set_cell_font(table.rows[r].cells[2], "就業時間", 6, True)
        time_str = f"{self._format_time(data.get('work_start_time'))}～{self._format_time(data.get('work_end_time'))}"
        self._set_cell_font(table.rows[r].cells[3], time_str, 6)
        self._set_cell_font(table.rows[r].cells[4], "休憩", 6, True)
        self._set_cell_font(table.rows[r].cells[5], f"{data.get('break_minutes', 60)}分", 6)

        r += 1
        # Row 7: Overtime
        self._set_cell_font(table.rows[r].cells[0], "時間外", 6, True)
        ot = f"1日{data.get('overtime_max_day', 4)}H/月{data.get('overtime_max_month', 45)}H以内"
        self._set_cell_font(table.rows[r].cells[1], ot, 6)
        self._set_cell_font(table.rows[r].cells[2], "休日労働", 6, True)
        self._set_cell_font(table.rows[r].cells[3], f"月{data.get('holiday_work_max', 2)}日以内", 6)
        self._set_cell_font(table.rows[r].cells[4], "深夜", 6, True)
        self._set_cell_font(table.rows[r].cells[5], "有（22:00-5:00）", 6)

        r += 1
        # Row 8: Rates (派遣料金 - for contract)
        self._set_cell_font(table.rows[r].cells[0], "派遣料金", 6, True)
        self._set_cell_font(table.rows[r].cells[1], f"¥{data.get('hourly_rate', 0):,.0f}/H", 6)
        self._set_cell_font(table.rows[r].cells[2], "時間外", 6, True)
        self._set_cell_font(table.rows[r].cells[3], f"¥{data.get('overtime_rate', 0):,.0f}", 6)
        self._set_cell_font(table.rows[r].cells[4], "休日", 6, True)
        self._set_cell_font(table.rows[r].cells[5], f"¥{data.get('holiday_rate', 0):,.0f}", 6)

        r += 1
        # Row 9: Wages (賃金 - for worker)
        self._set_cell_font(table.rows[r].cells[0], "賃金", 6, True)
        self._set_cell_font(table.rows[r].cells[1], f"¥{data.get('hourly_wage', 0):,.0f}/H", 6)
        self._set_cell_font(table.rows[r].cells[2], "締日", 6, True)
        self._set_cell_font(table.rows[r].cells[3], data.get('wage_closing', '月末'), 6)
        self._set_cell_font(table.rows[r].cells[4], "支払日", 6, True)
        self._set_cell_font(table.rows[r].cells[5], data.get('wage_payment', '翌25日'), 6)

        r += 1
        # Row 10: Haken-moto manager
        self._set_cell_font(table.rows[r].cells[0], "派遣元責任者", 6, True)
        moto_mgr = data.get('haken_moto_manager', {})
        self._set_cell_font(table.rows[r].cells[1], moto_mgr.get('name', ''), 6)
        self._set_cell_font(table.rows[r].cells[2], "TEL", 6, True)
        self._set_cell_font(table.rows[r].cells[3], moto_mgr.get('phone', ''), 6)
        self._set_cell_font(table.rows[r].cells[4], "部署", 6, True)
        self._set_cell_font(table.rows[r].cells[5], moto_mgr.get('department', ''), 6)

        r += 1
        # Row 11: Haken-saki manager
        self._set_cell_font(table.rows[r].cells[0], "派遣先責任者", 6, True)
        saki_mgr = data.get('haken_saki_manager', {})
        self._set_cell_font(table.rows[r].cells[1], saki_mgr.get('name', ''), 6)
        self._set_cell_font(table.rows[r].cells[2], "TEL", 6, True)
        self._set_cell_font(table.rows[r].cells[3], saki_mgr.get('phone', ''), 6)
        self._set_cell_font(table.rows[r].cells[4], "部署", 6, True)
        self._set_cell_font(table.rows[r].cells[5], saki_mgr.get('department', ''), 6)

        r += 1
        # Row 12: Complaint handlers
        self._set_cell_font(table.rows[r].cells[0], "苦情処理(元)", 6, True)
        moto_comp = data.get('haken_moto_complaint', {})
        self._set_cell_font(table.rows[r].cells[1], f"{moto_comp.get('name', '')} {moto_comp.get('phone', '')}", 5)
        self._set_cell_font(table.rows[r].cells[2], "苦情処理(先)", 6, True)
        saki_comp = data.get('haken_saki_complaint', {})
        table.rows[r].cells[3].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[3], f"{saki_comp.get('name', '')} {saki_comp.get('phone', '')}", 5)

        r += 1
        # Row 13: Safety & Termination
        self._set_cell_font(table.rows[r].cells[0], "安全衛生", 6, True)
        table.rows[r].cells[1].merge(table.rows[r].cells[2])
        self._set_cell_font(table.rows[r].cells[1], data.get('safety_measures', '派遣先規程に従う')[:20], 5)
        self._set_cell_font(table.rows[r].cells[3], "契約解除", 6, True)
        table.rows[r].cells[4].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[4], "30日前予告", 5)

        r += 1
        # Row 14: Welfare
        self._set_cell_font(table.rows[r].cells[0], "福利厚生", 6, True)
        welfare = data.get('welfare_facilities', [])
        if isinstance(welfare, list):
            welfare = '、'.join(welfare)
        table.rows[r].cells[1].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[1], welfare, 5)

        r += 1
        # Row 15: Insurance (for worker)
        self._set_cell_font(table.rows[r].cells[0], "社会保険", 6, True)
        ins = []
        if data.get('has_health_insurance', True):
            ins.append("健保")
        if data.get('has_pension', True):
            ins.append("年金")
        if data.get('has_employment_insurance', True):
            ins.append("雇保")
        self._set_cell_font(table.rows[r].cells[1], "・".join(ins), 6)
        self._set_cell_font(table.rows[r].cells[2], "雇用形態", 6, True)
        emp_type = "☑無期 ☐有期" if data.get('is_indefinite') else "☐無期 ☑有期"
        self._set_cell_font(table.rows[r].cells[3], emp_type, 6)
        self._set_cell_font(table.rows[r].cells[4], "協定対象", 6, True)
        kyotei = "☑対象" if data.get('is_kyotei_taisho', True) else "☐対象外"
        self._set_cell_font(table.rows[r].cells[5], kyotei, 6)

        r += 1
        # Row 16: 60over / mukeiko
        self._set_cell_font(table.rows[r].cells[0], "60歳以上", 6, True)
        over60 = "☑該当" if data.get('is_over_60') else "☐非該当"
        self._set_cell_font(table.rows[r].cells[1], over60, 6)
        self._set_cell_font(table.rows[r].cells[2], "無期/60歳限定", 6, True)
        mukeiko = "☑該当" if data.get('is_mukeiko_60over') else "☐非該当"
        self._set_cell_font(table.rows[r].cells[3], mukeiko, 6)
        self._set_cell_font(table.rows[r].cells[4], "個人抵触日", 6, True)
        self._set_cell_font(table.rows[r].cells[5], self._format_date_short(data.get('personal_conflict_date')), 6)

        r += 1
        # Row 17: Work scope change (2024 requirement)
        self._set_cell_font(table.rows[r].cells[0], "業務変更範囲", 6, True)
        table.rows[r].cells[1].merge(table.rows[r].cells[2])
        self._set_cell_font(table.rows[r].cells[1], data.get('work_change_scope', '会社の定める業務'), 5)
        self._set_cell_font(table.rows[r].cells[3], "就業場所変更", 6, True)
        table.rows[r].cells[4].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[4], data.get('location_change_scope', '会社の定める場所'), 5)

        r += 1
        # Row 18: Contract renewal
        self._set_cell_font(table.rows[r].cells[0], "更新の有無", 6, True)
        renewal = data.get('renewal_policy', '更新する場合がある')
        table.rows[r].cells[1].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[1], renewal, 5)

        r += 1
        # Row 19: Notes
        self._set_cell_font(table.rows[r].cells[0], "備考", 6, True)
        table.rows[r].cells[1].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[1], data.get('notes', ''), 5)

        # Signature section
        r += 1
        self._set_cell_font(table.rows[r].cells[0], "【甲】派遣元", 6, True)
        table.rows[r].cells[1].merge(table.rows[r].cells[2])
        self._set_cell_font(table.rows[r].cells[1], f"{settings.COMPANY_NAME}　　　　　　　　印", 6)
        self._set_cell_font(table.rows[r].cells[3], "【乙】派遣先", 6, True)
        table.rows[r].cells[4].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[4], f"{data.get('client_company_name', '')}　　　　　　印", 6)

        r += 1
        self._set_cell_font(table.rows[r].cells[0], "住所", 6)
        table.rows[r].cells[1].merge(table.rows[r].cells[2])
        self._set_cell_font(table.rows[r].cells[1], settings.COMPANY_ADDRESS, 5)
        self._set_cell_font(table.rows[r].cells[3], "住所", 6)
        table.rows[r].cells[4].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[4], data.get('client_address', ''), 5)

        r += 1
        # Worker acknowledgment
        self._set_cell_font(table.rows[r].cells[0], "派遣労働者", 6, True)
        self._set_cell_font(table.rows[r].cells[1], data.get('worker_name', ''), 6)
        self._set_cell_font(table.rows[r].cells[2], "上記内容確認", 6)
        self._set_cell_font(table.rows[r].cells[3], "署名:", 6)
        table.rows[r].cells[4].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[4], "＿＿＿＿＿＿＿＿＿", 6)

        r += 1
        # Footer
        self._set_cell_font(table.rows[r].cells[0], "作成日", 6)
        table.rows[r].cells[1].merge(table.rows[r].cells[5])
        self._set_cell_font(table.rows[r].cells[1], f"{self._format_date_japanese(date.today())}　※本書は労働者派遣法第26条・第34条に基づく", 5)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
